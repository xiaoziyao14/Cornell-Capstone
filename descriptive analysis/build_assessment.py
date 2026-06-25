"""
Builds dataset_clean_assessment.docx — explains how dataset_clean.csv differs from
the original Unicorn_Companies.csv and why, then gives a descriptive analysis of the
non-unicorn companies it adds. Neutral, factual tone. Generates its own charts (nu_*.png).
"""
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import warnings
warnings.filterwarnings('ignore')

DIR = "/Users/ziyaoxiao/Desktop/肖梓耀/Cornell/Capstone Project/descriptive analysis"
DATA = "/Users/ziyaoxiao/Desktop/肖梓耀/Cornell/Capstone Project/Data/dataset_clean.csv"
ACCENT = RGBColor(0x1D, 0x4E, 0xD8)

plt.rcParams.update({
    'font.family': 'DejaVu Sans', 'axes.spines.top': False, 'axes.spines.right': False,
    'axes.titlesize': 13, 'axes.titleweight': 'bold', 'figure.dpi': 150,
})
PAL = ['#2563EB','#7C3AED','#059669','#DC2626','#D97706','#0891B2',
       '#BE185D','#65A30D','#9333EA','#EA580C','#0D9488','#4F46E5','#6B7280']
U_COLOR, N_COLOR = '#2563EB', '#059669'

# ── load ────────────────────────────────────────────────────────────────────
df = pd.read_csv(DATA)
u = df[df['is_unicorn'] == 1].copy()
n = df[df['is_unicorn'] == 0].copy()
N = len(n)

# ════════════════════════════════════════════════════════════════════════════
# CHARTS (neutral titles)
# ════════════════════════════════════════════════════════════════════════════
# NU1 – geography
cont = n['Continent'].value_counts()
fig, ax = plt.subplots(figsize=(9, 5))
bars = ax.barh(cont.index[::-1], cont.values[::-1], color=PAL[:len(cont)], edgecolor='white')
for b, v in zip(bars, cont.values[::-1]):
    ax.text(b.get_width()+4, b.get_y()+b.get_height()/2, f'{v} ({v/N*100:.1f}%)',
            va='center', fontsize=10)
ax.set_xlim(0, cont.max()*1.25)
ax.set_xlabel('Number of Non-Unicorn Companies')
ax.set_title('Non-Unicorn Companies by Continent', fontsize=13, fontweight='bold')
plt.tight_layout(); plt.savefig(f'{DIR}/nu_01_geography.png', bbox_inches='tight'); plt.close()

# NU2 – industry
ind = n['Industry'].value_counts().head(10)
fig, ax = plt.subplots(figsize=(10, 6))
bars = ax.barh(ind.index[::-1], ind.values[::-1], color=PAL[:len(ind)], edgecolor='white')
for b, v in zip(bars, ind.values[::-1]):
    ax.text(b.get_width()+2, b.get_y()+b.get_height()/2, f'{v} ({v/N*100:.1f}%)',
            va='center', fontsize=9)
ax.set_xlim(0, ind.max()*1.25)
ax.set_xlabel('Number of Non-Unicorn Companies')
ax.set_title('Non-Unicorn Companies — Top 10 Industries', fontsize=13, fontweight='bold')
plt.tight_layout(); plt.savefig(f'{DIR}/nu_02_industry.png', bbox_inches='tight'); plt.close()

# NU3 – funding distribution
fb = [-1, 1e6, 10e6, 50e6, 100e6, 1e12]
fl = ['<$1M', '$1-10M', '$10-50M', '$50-100M', '$100M+']
fc = pd.cut(n['Funding_USD'], fb, labels=fl).value_counts().reindex(fl)
fig, ax = plt.subplots(figsize=(9, 5.5))
bars = ax.bar(fl, fc.values, color=PAL[:len(fl)], edgecolor='white')
for b, v in zip(bars, fc.values):
    ax.text(b.get_x()+b.get_width()/2, b.get_height()+5, f'{v}\n({v/N*100:.1f}%)',
            ha='center', fontsize=9)
ax.set_ylim(0, fc.max()*1.15)
ax.set_ylabel('Number of Companies'); ax.set_xlabel('Total Funding')
ax.set_title('Non-Unicorn Companies — Funding Distribution', fontsize=13, fontweight='bold')
plt.tight_layout(); plt.savefig(f'{DIR}/nu_03_funding.png', bbox_inches='tight'); plt.close()

# NU4 – year founded comparison (neutral)
yu = u['Year_Founded'].value_counts().sort_index()
yn = n['Year_Founded'].value_counts().sort_index()
fig, ax = plt.subplots(figsize=(11, 5.5))
ax.bar(yu.index-0.2, yu.values, width=0.4, color=U_COLOR, label='Unicorn', alpha=0.85)
ax.bar(yn.index+0.2, yn.values, width=0.4, color=N_COLOR, label='Non-Unicorn', alpha=0.85)
ax.set_xlabel('Year Founded'); ax.set_ylabel('Number of Companies')
ax.set_title('Year Founded — Unicorn vs Non-Unicorn', fontsize=13, fontweight='bold')
ax.legend(fontsize=11)
plt.tight_layout(); plt.savefig(f'{DIR}/nu_04_year_founded.png', bbox_inches='tight'); plt.close()

# NU5 – funding by class (neutral)
fb2 = [-1, 1e6, 10e6, 50e6, 100e6, 500e6, 1e12]
fl2 = ['<$1M','$1-10M','$10-50M','$50-100M','$100-500M','$500M+']
uc = pd.cut(u['Funding_USD'], fb2, labels=fl2).value_counts().reindex(fl2)
nc = pd.cut(n['Funding_USD'], fb2, labels=fl2).value_counts().reindex(fl2)
x = np.arange(len(fl2))
fig, ax = plt.subplots(figsize=(11, 5.5))
ax.bar(x-0.2, nc.values, width=0.4, color=N_COLOR, label='Non-Unicorn', alpha=0.85)
ax.bar(x+0.2, uc.values, width=0.4, color=U_COLOR, label='Unicorn', alpha=0.85)
ax.set_xticks(x); ax.set_xticklabels(fl2)
ax.set_ylabel('Number of Companies'); ax.set_xlabel('Total Funding')
ax.set_title('Funding by Class — Unicorn vs Non-Unicorn', fontsize=13, fontweight='bold')
ax.legend(fontsize=11)
plt.tight_layout(); plt.savefig(f'{DIR}/nu_05_funding_by_class.png', bbox_inches='tight'); plt.close()
print("Charts nu_01..nu_05 saved.")

# ════════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1); p.runs[0].font.color.rgb = ACCENT; return p
def h2(t): return doc.add_heading(t, level=2)
def para(t, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic; return p
def bullets(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''; c.paragraphs[0].add_run(hd).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = str(v)
    return t
def chart(f, w=6.0):
    doc.add_picture(f'{DIR}/{f}', width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
def meaning(t):
    p = doc.add_paragraph(); r = p.add_run('What this means — '); r.bold = True
    r.font.color.rgb = ACCENT; p.add_run(t)

# Title
title = doc.add_heading('dataset_clean.csv — Differences from the Original & Non-Unicorn Profile',
                        level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Cornell Capstone Project  •  dataset_clean.csv (n = 2,046: '
              f'{len(u)} unicorns + {N} non-unicorns)')
r.italic = True; r.font.size = Pt(10)
para('A teammate prepared dataset_clean.csv to support a supervised prediction task. This '
     'document first explains how it differs from the original Unicorn_Companies.csv and the '
     'reasoning behind those changes (Section 1), then profiles the non-unicorn companies it '
     'adds (Section 2). Each chart and table is followed by a short "What this means" note.',
     italic=True)

# ── SECTION 1 ───────────────────────────────────────────────────────────────
h1('Section 1 — How dataset_clean Differs from the Original (and Why)')
para('dataset_clean.csv is a purpose-built reshaping of the original data for a prediction '
     'task. The unicorn records themselves are unchanged — funding figures match the original '
     'to the dollar for 1,053 of 1,063 overlapping companies — so the changes are about '
     'structure and scope, not about altering unicorn facts.')

h2('1.1 Structure at a Glance')
table(['', 'Original (Unicorn_Companies)', 'dataset_clean'],
      [['Rows', '1,074 (all unicorns)', '2,046 (1,063 unicorns + 983 non-unicorns)'],
       ['Columns', '10', '9'],
       ['Label', 'none', 'is_unicorn (1/0)']])

h2('1.2 Columns Removed — and the Reasoning')
para('Three columns from the original were intentionally dropped. For a model that predicts '
     'whether a company is a unicorn, each of these would leak the answer or be unavailable '
     'for non-unicorns, so removing them is sound preparation:')
table(['Removed Column', 'Why It Was Removed'],
      [['Valuation', 'A unicorn is defined as valuation ≥ $1B, so it is effectively the '
        'answer itself — keeping it would be target leakage.'],
       ['Date Joined', 'The date a company became a unicorn exists only for unicorns; '
        'non-unicorns have no such date.'],
       ['Select Investors', 'Free-text investor lists are hard to use as a model feature and '
        'are not consistently available for smaller companies.']])
meaning('Removing these columns was a deliberate, reasonable choice for a prediction setup, '
        'not an oversight. The trade-off is that Date Joined and Select Investors were useful '
        'for descriptive work (they powered our timeline and investor analyses on the '
        'original data), so those angles are simply not available in dataset_clean.')

h2('1.3 Columns & Rows Added')
bullets([
    'is_unicorn — the 1/0 label that turns the data into a supervised learning problem.',
    'Description — a short text blurb per company (present for most non-unicorns; '
    'often blank for unicorns, which were carried over from the original).',
    'Funding_USD — the original "$8B"-style funding strings converted to clean numeric USD.',
    '983 non-unicorn companies — the negative class needed for supervised learning.',
    'Note: 9 pre-1995 unicorns from the original (e.g. Otto Bock HealthCare, founded 1919) '
    'were not carried over, so both classes now start in 1995 — a reasonable trimming of '
    'non-startup legacy firms.',
])
meaning('In short, dataset_clean keeps the original unicorn facts, drops fields that do not '
        'fit a prediction task, converts funding to a usable numeric form, and adds a '
        'negative class. The genuinely new material is the 983 non-unicorn companies, '
        'profiled next.')

# ── SECTION 2 ───────────────────────────────────────────────────────────────
doc.add_page_break()
h1('Section 2 — Descriptive Analysis of the Non-Unicorn Companies')
para(f'The {N} non-unicorn companies are private firms that had not reached $1B valuation. '
     'This section describes who they are and how they compare with the unicorn group.')

h2('2.1 Geographic Distribution')
table(['Continent', 'Companies', 'Share'],
      [[k, v, f'{v/N*100:.1f}%'] for k, v in n['Continent'].value_counts().items()])
para('Top countries: ' + ', '.join(f'{k} ({v})' for k, v in n['Country'].value_counts().head(5).items()) +
     '. (96 companies are recorded as continent "Other" and 72 have no country listed.)')
chart('nu_01_geography.png', 5.6)
meaning('The non-unicorn set leans more heavily toward North America (65%) than the unicorn '
        'population does (55%). It reflects wherever this particular list was sourced, so it '
        'is best read as a sample of companies rather than a global census.')

h2('2.2 Industry Distribution')
table(['Industry', 'Companies', 'Share'],
      [[k, v, f'{v/N*100:.1f}%'] for k, v in n['Industry'].value_counts().head(6).items()])
chart('nu_02_industry.png', 5.8)
meaning('Internet Software & Services is the largest category at 37%, an even heavier share '
        'than among unicorns (19%). The negative class is concentrated in software companies, '
        'which is useful context when comparing the two groups by sector.')

h2('2.3 Funding Distribution')
para(f'Median funding among non-unicorns is ${n["Funding_USD"].median()/1e6:.2f}M, compared '
     f'with ${u["Funding_USD"].median()/1e6:.0f}M for unicorns.')
table(['Funding Band', 'Companies', 'Share'],
      [['< $1M', 565, '57.5%'], ['$1-10M', 300, '30.5%'], ['$10-50M', 103, '10.5%'],
       ['$50-100M', 8, '0.8%'], ['$100M+', 7, '0.7%']])
chart('nu_03_funding.png', 5.6)
meaning('The non-unicorns are mostly earlier-stage or smaller companies — 88% raised under '
        '$10M. This is the clearest structural difference between the two groups: unicorns '
        'cluster at much larger funding levels.')

h2('2.4 Funding by Class')
chart('nu_05_funding_by_class.png', 6.0)
meaning('Placed side by side, the two groups occupy largely different funding ranges, with '
        'limited overlap in the middle. Funding is therefore the dimension that most clearly '
        'distinguishes the classes in this dataset.')

h2('2.5 Founding Year')
para(f'Non-unicorn founding years range {int(n["Year_Founded"].min())}–'
     f'{int(n["Year_Founded"].max())} (median {int(n["Year_Founded"].median())}); the count '
     'peaks around 2013. The unicorn group extends further, through 2021.')
chart('nu_04_year_founded.png', 6.2)
meaning('The two groups cover somewhat different time windows — the non-unicorns were '
        'collected mostly up to the early 2010s, while unicorns continue to 2021. Worth '
        'keeping in mind when comparing the groups, so that founding year is not mistaken for '
        'a difference in company quality.')

# ── CLOSING NOTE ────────────────────────────────────────────────────────────
h1('Closing Note')
para('dataset_clean.csv is a sensible, prediction-oriented reshaping of the original data: '
     'unicorn facts are preserved, leakage-prone columns were removed on purpose, funding was '
     'made numeric, and a non-unicorn class was added. The main practical differences between '
     'the two groups — funding scale and time window — are natural consequences of how each '
     'set was assembled, and are simply useful to keep in mind for any later modeling.')

out = f'{DIR}/dataset_clean_assessment.docx'
doc.save(out)
print('Saved:', out)
