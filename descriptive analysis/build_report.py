"""
Builds the Word report (Unicorn_Descriptive_Analysis_Report.docx) that
compiles both the guided analysis and the independent data-driven insights,
with tables and embedded charts.
"""
import pandas as pd
import numpy as np
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DIR = "/Users/ziyaoxiao/Desktop/肖梓耀/Cornell/Capstone Project/descriptive analysis"
DATA = "/Users/ziyaoxiao/Desktop/肖梓耀/Cornell/Capstone Project/Data/Unicorn_Companies.csv"
ACCENT = RGBColor(0x1D, 0x4E, 0xD8)

# ── load & clean (mirrors the analysis scripts) ─────────────────────────────
df = pd.read_csv(DATA)
canon = {}
for n in df['Industry'].unique():
    k = n.lower()
    if k not in canon:
        canon[k] = n
df['Industry'] = df['Industry'].apply(lambda s: canon[s.lower()])

def pd_(s):
    s = str(s).replace('$', '').replace(',', '')
    if s.endswith('B'): return float(s[:-1]) * 1e9
    if s.endswith('M'): return float(s[:-1]) * 1e6
    try: return float(s)
    except: return np.nan

df['V'] = df['Valuation'].apply(pd_)
df['F'] = df['Funding'].apply(pd_)
df['YJ'] = pd.to_datetime(df['Date Joined']).dt.year
df['YTU'] = df['YJ'] - df['Year Founded']
df['eff'] = df['V'] / df['F']
total = len(df)

# ── document helpers ────────────────────────────────────────────────────────
doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = ACCENT
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold
    return p

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        run = c.paragraphs[0].add_run(hd)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t

def add_chart(fname, width=6.3):
    doc.add_picture(f'{DIR}/{fname}', width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def meaning(text):
    """Interpretation paragraph: bold 'What this means' lead-in + takeaway."""
    p = doc.add_paragraph()
    lead = p.add_run('What this means — ')
    lead.bold = True
    lead.font.color.rgb = ACCENT
    p.add_run(text)
    return p

# ════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('Unicorn Companies — Descriptive Analysis Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Cornell Capstone Project  •  Dataset: Unicorn_Companies.csv (n = 1,074)  •  '
                'Coverage through 2022-04-05')
r.italic = True; r.font.size = Pt(10)

para('This report describes the population of existing unicorn companies '
     '(private, venture-backed startups valued at ≥ $1B). Section A presents the '
     'guided descriptive analysis; Section B presents an independent, data-driven '
     'insight scan; Section C notes implications for modeling. Every chart and table is '
     'followed by a "What this means" note stating the takeaway, so the focus is on what '
     'the data implies rather than merely what was computed. All per-category shares '
     'describe this dataset’s composition, not real-world probabilities.', italic=True)

# Data cleaning note
h2('Data Preparation')
bullets([
    'Industry capitalization inconsistency merged (Artificial Intelligence × 11 + '
    'Artificial intelligence × 73 → 84 companies, 7.8%).',
    'Valuation and Funding parsed from strings (e.g. "$180B") into numeric USD.',
    'Date Joined parsed to datetime; derived Year Joined and Years-to-Unicorn.',
])

# ════════════════════════════════════════════════════════════════════════════
# SECTION A – GUIDED ANALYSIS
# ════════════════════════════════════════════════════════════════════════════
h1('Section A — Guided Descriptive Analysis')
para('This section follows the analytical questions defined by the team: where unicorns '
     'are located, what industries they occupy, how they are valued and funded, and how '
     'their industry mix has shifted across technological eras.')

# A1 Geography
h2('A1. Geographic Distribution')
cont = df['Continent'].value_counts()
add_table(['Continent', 'Unicorns', 'Share'],
          [[k, v, f'{v/total*100:.1f}%'] for k, v in cont.items()])
top_ctry = df['Country'].value_counts().head(5)
para('Top countries: ' + ', '.join(f'{k} ({v})' for k, v in top_ctry.items()) +
     '. The US, China, and India together account for 74.5% of all unicorns.')
add_chart('01_geographic_distribution.png')
meaning('Unicorn creation is highly concentrated, not globally spread. Three countries '
        'produce three-quarters of all unicorns, which tells us that breakout outcomes '
        'depend less on market size and more on the maturity of the surrounding ecosystem '
        '— deep venture capital, technical talent, and viable exit markets. For a '
        'prediction task, a company’s location is therefore a meaningful signal, not '
        'just a label.')

# A2 Industry
h2('A2. Industry Distribution')
ind = df['Industry'].value_counts().head(8)
add_table(['Industry', 'Unicorns', 'Share'],
          [[k, v, f'{v/total*100:.1f}%'] for k, v in ind.items()])
para('Fintech (20.9%) and Internet Software & Services (19.1%) together make up ~40% '
     'of all unicorns.')
add_chart('02_industry_distribution.png')
meaning('Value creation clusters in software-driven, digitally scalable business models. '
        'The two leading categories — Fintech and Internet Software — are both built on '
        'code that scales at near-zero marginal cost, which is exactly what lets a company '
        'reach a billion-dollar valuation quickly. Capital-intensive, physical-asset '
        'sectors sit far down the list. The takeaway: industry is one of the strongest '
        'descriptors of where unicorns come from.')

# A3 Valuation & Funding
h2('A3. Valuation & Funding')
bullets([
    f'Median valuation ${df["V"].median()/1e9:.0f}B; 94.3% sit in the $1B–$10B range; '
    'ByteDance is the extreme top at $180B.',
    f'Median total funding ${df["F"].median()/1e6:.0f}M; 65.4% raised under $500M.',
    'Valuation-to-funding ratio is roughly 5:1, reflecting investor growth premiums.',
])
add_chart('03_valuation_distribution.png', width=5.6)
add_chart('04_funding_distribution.png', width=5.6)
meaning('Most unicorns are "just barely" unicorns. About 84% sit in the $1B–$5B band, '
        'clustered right above the $1B threshold, while true giants are vanishingly rare. '
        'The funding side tells a complementary story: companies are valued at roughly 5x '
        'the cash they actually raised, which means the market is paying a large premium '
        'for expected future growth rather than for money already invested. Valuation is '
        'driven by narrative and growth potential, not by capital alone.')

# A4 Time to unicorn
h2('A4. Time to Unicorn')
y = df[df['YTU'] >= 0]
para(f'Median time from founding to unicorn status is {y["YTU"].median():.0f} years '
     f'(mean {y["YTU"].mean():.1f}). The distribution is wide, with rare extreme outliers.')
add_chart('06_years_to_unicorn.png', width=5.8)
meaning('Among companies that did become unicorns, the typical one took 6–7 years to '
        'reach $1B — so this is a medium-term outcome rather than an overnight one. '
        'Important caveat: because the dataset contains unicorns only, this chart cannot '
        'be read as "taking longer raises the chance of becoming a unicorn." We have no '
        'data on the companies that never reached $1B, so any statement about the '
        'probability of becoming a unicorn would be survivorship bias. The chart describes '
        'the timing of successes, not the odds of success. The long right tail also flags '
        'a data-quality caveat: a handful of decades-old firms appear here and are not '
        'typical venture startups, which matters when this variable is used as a feature.')

# A5 Eras (the human-judgment view)
h2('A5. Technological Eras (Guided Hypothesis)')
para('Framing the data around the dominant technology wave of each period: companies are '
     'grouped by founding year into four eras. This view tests the hypothesis that each '
     'technological shift produced a distinct cohort of breakout companies.')
add_table(['Era', 'Founded', 'Unicorns', 'Dominant Wave'],
          [['Dot-Com & Early Web', '≤2000', '37 (3.4%)', 'First commercial internet'],
           ['Web 2.0 & Social', '2001–2008', '109 (10.1%)', 'Broadband + social media'],
           ['Mobile & Cloud', '2009–2014', '447 (41.6%)', 'Smartphones, cloud, sharing economy'],
           ['AI & Modern Stack', '2015–2021', '481 (44.8%)', 'AI/ML mainstream, modern fintech']])
bullets([
    'Internet Software & Services peaked in the Web 2.0 era (25.7%) and has declined since '
    'to 14.6% as the web commoditized.',
    'Fintech rose continuously across every era — 11.9% → 17.4% → 26.2% — becoming the '
    'single largest category in the modern stack.',
    'Artificial Intelligence was 0% before 2009, then appeared with the Mobile & Cloud era '
    '(8.3%) and held ~9% afterward.',
])
add_chart('08_technological_eras.png')
meaning('Each technology wave left a visible fingerprint on which industries produced '
        'unicorns: the Web 2.0 era was an internet-software story, while the modern stack '
        'is a Fintech-and-AI story. The data tells a clear directional narrative — old '
        'categories fade and new ones rise as the underlying technology shifts.')
para('Statistical footnote: a chi-square test confirms the Era–Industry association is '
     'statistically significant (p = 0.0003) but weak in magnitude (Cramér’s V = 0.16). '
     'In plain terms, the industry mix genuinely shifts across eras, yet era alone '
     'explains only a small part of it — geography is a stronger structural driver '
     '(Country–Industry V = 0.25; Continent–Industry V = 0.22). Eras are defined by '
     'founding year, and because the data ends in April 2022 the most recent AI surge '
     '(e.g. post-ChatGPT) is not captured.', italic=True)

# ════════════════════════════════════════════════════════════════════════════
# SECTION B – INDEPENDENT INSIGHTS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('Section B — Independent Data-Driven Insights')
para('This section was generated without predefined hypotheses, scanning the data '
     '(including the previously unused Select Investors field) for non-obvious patterns.')

# B1 Capital efficiency
h2('B1. Capital Efficiency Varies Sharply by Industry')
valid = df[(df['F'] > 0) & df['F'].notna() & df['V'].notna()]
para(f'Overall median valuation-to-funding multiple is {valid["eff"].median():.1f}x. '
     'Software-type sectors convert capital into valuation far more efficiently than '
     'asset-heavy sectors.')
add_table(['Most Efficient', 'x', 'Least Efficient', 'x'],
          [['Internet Software', '6.4x', 'Auto & Transportation', '2.7x'],
           ['Fintech', '6.2x', 'Travel', '2.8x'],
           ['Data Mgmt & Analytics', '6.2x', 'Supply Chain & Logistics', '3.5x']])
add_chart('09_capital_efficiency.png', width=5.8)
meaning('Not all unicorns are built the same way. Software sectors turn each dollar raised '
        'into more than twice the valuation of asset-heavy sectors like transportation and '
        'travel, which must keep burning cash on physical operations. This makes capital '
        'efficiency a meaningful way to separate "asset-light" from "asset-heavy" unicorns '
        '— a distinction the raw industry label alone does not capture, and a strong '
        'candidate engineered feature for clustering.')

# B2 Investors
h2('B2. A Small Set of Investors Backs a Large Share')
inv = Counter()
for s in df['Select Investors'].dropna():
    for i in s.split(','):
        inv[i.strip()] += 1
add_table(['Investor', 'Unicorns Backed'],
          [[k, v] for k, v in inv.most_common(8)])
para('Accel, Tiger Global, Andreessen Horowitz, and the Sequoia family each appear in '
     '~50–60 unicorns — a concentrated backer network behind the population.')
meaning('The same handful of elite investors recur behind a large slice of all unicorns. '
        'This suggests that who funds a company carries real signal: backing from a '
        'top-tier fund is both a vote of confidence and an accelerant (capital, network, '
        'credibility). The previously unused Select Investors field could be turned into '
        'a powerful feature — e.g. a flag for "backed by a prolific unicorn-maker."')

# B3 City concentration
h2('B3. Unicorns Are a City Phenomenon, Not Just a Country One')
city = df['City'].value_counts().head(6)
add_table(['City', 'Unicorns', 'Share'],
          [[k, v, f'{v/total*100:.1f}%'] for k, v in city.items()])
para('The top 3 cities (San Francisco, New York, Beijing) alone account for 29.6% of all '
     'unicorns; San Francisco (14.2%) exceeds half of all of Europe.')
meaning('Unicorns concentrate at the city level even more tightly than at the country '
        'level — they cluster in a few dense hubs where capital, talent, and founders '
        'physically collide. This means geography matters at a finer grain than '
        '"continent" or even "country"; the specific city carries information that the '
        'broader region washes out.')

# B4 National specialization
h2('B4. Strong National Industry Specialization')
para('Countries are not industry-neutral — several concentrate heavily in one sector, '
     'meaning Country and Industry are correlated rather than independent.')
add_table(['Country', '#1 Industry', 'Share of Country’s Unicorns'],
          [['United Kingdom', 'Fintech', '60%'],
           ['Israel', 'Cybersecurity', '30%'],
           ['India', 'E-commerce', '25%'],
           ['United States', 'Internet Software', '27%'],
           ['China', 'E-commerce', '17%']])
add_chart('10_country_specialization.png', width=5.8)
meaning('Countries are not neutral mixes of every industry — each ecosystem specializes, '
        'reflecting local strengths (UK finance, Israeli defense-born cybersecurity, '
        'India’s consumer-internet scale). The practical consequence is that Country and '
        'Industry carry overlapping information rather than independent information, so a '
        'model should not treat them as unrelated features or it will double-count the '
        'same underlying signal.')

# B5 Speed extremes
h2('B5. Two Extremes: Instant Unicorns and Late Bloomers')
fast = y[y['YTU'] <= 2]
old = df[df['YTU'] > 20]
bullets([
    f'{len(fast)} companies ({len(fast)/len(y)*100:.1f}%) reached $1B within 2 years of '
    'founding — 57 of them in 2021 alone, evidence of the capital bubble.',
    f'{len(old)} companies took more than 20 years; the extreme case, Otto Bock HealthCare, '
    'was founded in 1919 (98 years).',
])
para('Implication: not all rows are true startups. Aged outliers contaminate the '
     '"startup" assumption and should be handled before modeling.', italic=True)
meaning('The population hides two very different stories at its extremes. The "instant '
        'unicorns" cluster tightly in 2021, which exposes how much the boom was a '
        'function of cheap capital and timing rather than company fundamentals alone. The '
        '"late bloomers" reveal a data-cleanliness problem: a 1919-founded firm is not a '
        'venture startup, so the dataset mixes genuine startups with re-labeled legacy '
        'companies — these outliers must be capped or filtered before any distance-based '
        'modeling.')

# B6 Funding-valuation weak link
h2('B6. Funding and Valuation Are Only Weakly Linked')
sc = df[['V','F']].dropna(); sc = sc[(sc['F'] > 0) & (sc['V'] > 0)]
corr = np.log(sc['V']).corr(np.log(sc['F']))
para(f'The log-log correlation between funding and valuation is only {corr:.2f}, meaning '
     f'money raised explains roughly {corr**2*100:.0f}% of valuation variance. Counter to '
     'intuition, "more money raised" does not strongly imply "higher valuation" — sector, '
     'geography, and timing dominate.')
add_chart('11_funding_vs_valuation.png', width=5.6)
meaning('This is the most important caution for modeling. The intuitive assumption — '
        '"companies that raise more are worth more" — is only weakly true here: funding '
        'accounts for about a third of the variation in valuation, and the scatter is '
        'visibly wide. A model that leans on funding as its main predictor would miss most '
        'of the story; the remaining ~64% lives in industry, geography, timing, and '
        'investor quality. Rich, multi-dimensional features are required.')

# ════════════════════════════════════════════════════════════════════════════
# SECTION C – IMPLICATIONS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('Section C — Implications for Modeling')
bullets([
    'Coverage ends April 2022 — no data on 2023+ unicorns; recent-trend questions need a newer source.',
    'This dataset contains unicorns only, so it cannot by itself train a model that predicts '
    'whether a company becomes a unicorn; a labeled set with non-unicorns is required '
    '(see Data/dataset_clean.csv).',
    'Right-skewed Valuation/Funding and outliers in Years-to-Unicorn should be transformed '
    'or capped before distance-based methods such as K-means.',
    'Country and Industry are correlated (national specialization) — avoid treating them as independent.',
    'Capital efficiency (Valuation/Funding) is a useful engineered feature that separates '
    'asset-light from asset-heavy unicorns.',
    'Funding alone is a weak predictor of valuation (r ≈ 0.60); richer features are needed.',
])

out_path = f'{DIR}/Unicorn_Descriptive_Analysis_Report.docx'
doc.save(out_path)
print('Saved:', out_path)
